#!/usr/bin/env python
# encoding: utf-8

"""


##################
#
#	views for exporting data
#
##################


"""

from django.http import HttpResponse, Http404, HttpResponseRedirect
from django.shortcuts import render_to_response, redirect, get_object_or_404
from django.template import RequestContext
# from django.utils.html import strip_tags, escape
from django.utils.encoding import smart_str
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator, InvalidPage, EmptyPage
from django.core.urlresolvers import reverse

import datetime
import csv


from koncepts.models import *
from koncepts.export import *
from settings import printdebug, MEDIA_ROOT


# ======
# ===========
# WORD
# ===========
# ======



try:
	from docx import Document
	from docx.shared import Inches
except:
	print "Error importing docx library!"


@login_required
def export_word(request):
	"""
	Export data in DOCX format

	http://127.0.0.1:8000/api/export/word?quote_id=100

	info:
	http://python-docx.readthedocs.org/en/latest/user/styles-understanding.html#understanding-styles

	"""

	# try to extract a valid instance
	valid_obj = get_object(request)

	if valid_obj:

		response = HttpResponse(mimetype='text/docx')
		response['Content-Disposition'] = 'attachment; filename=' + valid_obj['filename'] + '.docx'

		wordDocument = Document()


		# ---
		# contents
		# ---

		if 'quote' in valid_obj.keys():

			quote = valid_obj['quote']

			wordDocument = writeSingleQuote(wordDocument, request, quote)

		elif 'document' in valid_obj.keys():

			document = valid_obj['document']

			for quote in document.fragment_set.all():
				wordDocument = writeSingleQuote(wordDocument, request, quote)


		elif 'koncept' in valid_obj.keys():

			koncept = valid_obj['koncept']

			for intf in koncept.intfrag_set.all():
				wordDocument = writeSingleQuote(wordDocument, request, intf.fragment)

		else:
			raise Http404


		wordDocument.save(response)
		return response


	else:
		raise Http404







# ---
# single quote
# ---

def writeSingleQuote(wordDocument, request, quote):
	""" writes a single item to the document """
	uri = request.build_absolute_uri(quote.get_absolute_url())
	created = "Created on: {:%d, %b %Y}".format(quote.created_at)
	downloaded = "Downloaded on: {:%d, %b %Y}".format(datetime.date.today())

	if quote.source:
		sourcetitle = quote.source.title
		sourceauthor = quote.source.author
		sourcepubyear = quote.source.pubyear or ""
		sourcedesc = quote.source.description
		sourceurl = quote.source.url

	# compose the document

	# wordDocument.add_heading(uri, 1)
	wordDocument.add_heading(quote.title, 1)

	if True:
		wordDocument.add_paragraph("---------------------------------")
		p = wordDocument.add_paragraph()
		p.add_run(uri, 'SubtleEmphasis')
		p = wordDocument.add_paragraph()
		p.add_run(created, 'SubtleEmphasis')
		p = wordDocument.add_paragraph()
		p.add_run(downloaded, 'SubtleEmphasis')
		wordDocument.add_paragraph("---------------------------------")

	wordDocument.add_paragraph("")
	p = wordDocument.add_paragraph(style='Quote')
	p.add_run(quote.text, "Body Text Char")

	if quote.source:
		wordDocument.add_paragraph("")
		bigtitle = " ".join([sourcetitle, sourceauthor, str(sourcepubyear)])
		p = wordDocument.add_paragraph(style='List Paragraph')
		p.add_run(bigtitle, 'Book Title')

		if sourcedesc:
			wordDocument.add_paragraph(sourcedesc, style='List Paragraph')
		if sourceurl:
			p = wordDocument.add_paragraph(style='List Paragraph')
			p.add_run(quote.source.url, 'Emphasis')


	wordDocument.add_paragraph("")
	wordDocument.add_paragraph("")
	if False:
		p = wordDocument.add_paragraph()
		p.add_run(uri, 'SubtleEmphasis')
		p = wordDocument.add_paragraph()
		p.add_run(created, 'SubtleEmphasis')
		p = wordDocument.add_paragraph()
		p.add_run(downloaded, 'SubtleEmphasis')
		wordDocument.add_paragraph("---------------------------------")
		wordDocument.add_paragraph("")

	return wordDocument
